"""
从 Word / Excel 抽取纯文本，供「按说明补全」类大模型处理。
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import BinaryIO, Tuple, Union


def _require_docx():
    try:
        import docx

        return docx
    except ImportError as e:
        raise RuntimeError(
            "未安装 python-docx，请在项目目录执行: pip install python-docx"
        ) from e


def _require_openpyxl():
    try:
        import openpyxl

        return openpyxl
    except ImportError as e:
        raise RuntimeError(
            "未安装 openpyxl，请在项目目录执行: pip install openpyxl"
        ) from e


def extract_docx_text(data: Union[bytes, BinaryIO]) -> str:
    docx = _require_docx()
    bio = data if hasattr(data, "read") else io.BytesIO(data)
    d = docx.Document(bio)
    parts: list[str] = []
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in d.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n\n".join(parts).strip()


def extract_xlsx_text(data: Union[bytes, BinaryIO]) -> str:
    openpyxl = _require_openpyxl()
    bio = data if hasattr(data, "read") else io.BytesIO(data)
    wb = openpyxl.load_workbook(bio, read_only=True, data_only=True)
    try:
        out: list[str] = []
        for sn in wb.sheetnames:
            ws = wb[sn]
            out.append(f"=== 工作表：{sn} ===")
            for row in ws.iter_rows(values_only=True):
                line = "\t".join("" if v is None else str(v).strip() for v in row)
                if line.strip():
                    out.append(line)
        return "\n".join(out).strip()
    finally:
        wb.close()


def sniff_office_kind(filename: str) -> str:
    """返回 docx | xlsx | ''"""
    suf = Path(filename or "").suffix.lower()
    if suf == ".docx":
        return "docx"
    if suf in (".xlsx", ".xlsm"):
        return "xlsx"
    return ""


def extract_office_text(filename: str, data: bytes) -> Tuple[str, str]:
    """
    根据扩展名抽取正文。返回 (kind, text)，kind 为 docx/xlsx；
    不支持的扩展名抛 ValueError。
    """
    kind = sniff_office_kind(filename)
    if kind == "docx":
        return kind, extract_docx_text(data)
    if kind == "xlsx":
        return kind, extract_xlsx_text(data)
    raise ValueError("仅支持 .docx 与 .xlsx/.xlsm")
