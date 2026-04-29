"""
AI 文稿去痕：去除异常符号与空白，可选调用 OpenAI 兼容 API；导出 PDF / Word / Excel。
"""

from __future__ import annotations

import hashlib
import io
import logging
import os
import re
from typing import Any, Iterator, Optional, Tuple

from openai import OpenAI

from utils.docx_format_spec import generate_file_fill_system_prompt
from utils.openai_compat import get_openai_compat_config
from utils.text_format import normalize_display_text, sanitize_xml_compatible_text

logger = logging.getLogger(__name__)

_MAX_INPUT_CHARS = int(os.getenv("AI_CLEAN_MAX_CHARS", "200000"))


def _patch_hashlib_md5_for_reportlab() -> None:
    """
    ReportLab 在 pdfdoc 中调用 hashlib.md5(usedforsecurity=False)。
    部分环境（如 OpenSSL 后端）的 md5 不接受该参数，会触发 TypeError。
    """
    if getattr(hashlib, "_transvsverter_md5_patched", False):
        return
    _orig = hashlib.md5

    def md5(data=b"", *args, **kwargs):
        kwargs.pop("usedforsecurity", None)
        if args:
            return _orig(data, *args)
        return _orig(data)

    hashlib.md5 = md5  # type: ignore[assignment]
    setattr(hashlib, "_transvsverter_md5_patched", True)


_MAX_API_CHUNK = int(os.getenv("AI_CLEAN_CHUNK_CHARS", "12000"))


def _strip_markdown_artifacts(s: str) -> str:
    t = s
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"__([^_]+)__", r"\1", t)
    t = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", t)
    t = re.sub(r"(?<!_)_([^_\n]+)_(?!_)", r"\1", t)
    # 保留行首 # / ## / ### 标题结构（由生成与清理流程约定），不再整行去掉井号
    t = re.sub(r"^\s*[-*+]\s+", "", t, flags=re.MULTILINE)
    t = re.sub(r"^\s*\d+\.\s+", "", t, flags=re.MULTILINE)
    return t


def clean_text_local(raw: Optional[str]) -> str:
    if not raw:
        return ""
    t = _strip_markdown_artifacts(str(raw))
    t = normalize_display_text(t, multiline=True)
    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


_HEADING_LINE = re.compile(r"^(#{1,6})\s+(.+)$")


def iter_document_blocks(text: str) -> Iterator[Tuple[str, Optional[int], str]]:
    """
    解析正文为结构化块，供 PDF / Word 与前端预览使用。
    返回 (kind, level, content)：kind 为 'h'（level 1–3）或 'p'（level 为 None）。
    支持行首 Markdown 标题；无标题时按空行分段，段内换行压成空格以适配中文排版。
    """
    text = (text or "").strip()
    if not text:
        return
    lines = text.split("\n")
    has_heading = any(_HEADING_LINE.match(x.strip()) for x in lines)
    if not has_heading:
        parts = re.split(r"\n\s*\n+", text)
        if len(parts) == 1 and parts[0] and "\n" in parts[0]:
            alt = re.split(r"(?<=[。！？])\s*\n", parts[0])
            if len(alt) > 1:
                parts = [p.strip() for p in alt if p.strip()]
        for part in parts:
            part = part.strip()
            if not part:
                continue
            body = " ".join(part.split())
            if body:
                yield ("p", None, body)
        return
    i, n = 0, len(lines)
    while i < n:
        st = lines[i].strip()
        if not st:
            i += 1
            continue
        mh = _HEADING_LINE.match(st)
        if mh:
            level = min(len(mh.group(1)), 3)
            yield ("h", level, mh.group(2).strip())
            i += 1
            continue
        buf = [st]
        i += 1
        while i < n:
            s2 = lines[i].strip()
            if not s2:
                break
            if _HEADING_LINE.match(s2):
                break
            buf.append(s2)
            i += 1
        merged = "\n".join(buf)
        body = " ".join(merged.split())
        if body:
            yield ("p", None, body)


def _chunk_text_for_api(text: str, max_chunk: int) -> list[str]:
    if len(text) <= max_chunk:
        return [text]
    parts: list[str] = []
    buf: list[str] = []
    n = 0
    for para in text.split("\n"):
        plen = len(para) + (1 if buf else 0)
        if n + plen > max_chunk and buf:
            parts.append("\n".join(buf))
            buf = [para]
            n = len(para)
        else:
            buf.append(para)
            n += plen
    if buf:
        parts.append("\n".join(buf))
    return parts


def _extract_chat_completion_text(response: Any) -> str:
    """
    从 OpenAI 兼容 chat.completions 响应中取出 assistant 文本。
    ModelScope 等接口在 content 过滤、长度或异常时可能返回 choices 为空或 content 为 None，需避免崩溃。
    """
    try:
        choices = getattr(response, "choices", None)
        if not choices:
            logger.warning(
                "Chat completion：choices 为空 model=%s id=%s",
                getattr(response, "model", None),
                getattr(response, "id", None),
            )
            return ""
        ch0 = choices[0]
        if ch0 is None:
            logger.warning("Chat completion：choices[0] 为 None")
            return ""
        msg = getattr(ch0, "message", None)
        if msg is None:
            logger.warning(
                "Chat completion：message 为空 finish_reason=%s",
                getattr(ch0, "finish_reason", None),
            )
            return ""
        raw = getattr(msg, "content", None)
        if raw is None:
            logger.warning(
                "Chat completion：content 为空 finish_reason=%s refusal=%s",
                getattr(ch0, "finish_reason", None),
                getattr(msg, "refusal", None),
            )
            return ""
        if isinstance(raw, list):
            parts: list[str] = []
            for p in raw:
                if isinstance(p, dict) and p.get("type") == "text" and p.get("text") is not None:
                    parts.append(str(p["text"]))
                elif isinstance(p, str):
                    parts.append(p)
            raw = "".join(parts)
        return str(raw).strip()
    except (TypeError, IndexError, AttributeError) as e:
        logger.warning("Chat completion：解析响应异常 %s", e)
        return ""


# 接口可能 HTTP 200 但 choices/content 为空（策略拦截、模型异常、请求过长等）
EMPTY_CHAT_COMPLETION_USER_MESSAGE = (
    "大模型接口返回了 HTTP 200，但回复正文为空（choices 或 content 为空），无法完成编辑。"
    "可能被内容安全策略拦截、模型异常或单次请求过长，请缩短说明与文档、更换模型或稍后重试。"
)


_CLEAN_PROMPT = """你是文本编辑。请清理以下由 AI 或网页粘贴的文稿：
1. 保留层级结构：单独成行且以「# 」「## 」「### 」（井号+空格+标题文字）表示章节标题的，必须原样保留该行格式，不要删除井号或把标题并入正文
2. 去掉行内装饰：可删除 **、*、行内反引号等加粗/斜体符号；列表若可改为连贯段落则改，但不要破坏上述标题行
3. 正文按意群分段，段与段之间空一行；修正异常空格、全角半角混用、零宽字符与多余空行
4. 不要添加前言、标题「清理后」或任何解释，保持原意与信息完整，中文标点为主"""


def clean_text_with_api(
    client: OpenAI,
    model: str,
    raw: str,
) -> str:
    if not raw.strip():
        return ""
    chunks = _chunk_text_for_api(raw, _MAX_API_CHUNK)
    out_parts: list[str] = []
    for i, chunk in enumerate(chunks):
        user_msg = f"以下为第 {i + 1}/{len(chunks)} 段，请按说明清理：\n\n{chunk}"
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": _CLEAN_PROMPT},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.2,
            max_tokens=min(16384, len(chunk) * 2 + 2048),
        )
        part = _extract_chat_completion_text(response)
        out_parts.append(part)
    merged = "\n\n".join(p for p in out_parts if p)
    return normalize_display_text(merged, multiline=True)


def build_openai_client() -> Tuple[Optional[OpenAI], str]:
    cfg = get_openai_compat_config()
    if not cfg.api_key:
        return None, cfg.base_url
    try:
        timeout_sec = float(os.getenv("AI_CLEAN_OPENAI_TIMEOUT", "180"))
        # 429 时 SDK 默认会重试，易在短时间内多次命中限流；默认不重试，可用环境变量调回
        max_retries = int(os.getenv("AI_CLEAN_OPENAI_MAX_RETRIES", "0"))
        return OpenAI(
            base_url=cfg.base_url,
            api_key=cfg.api_key,
            timeout=timeout_sec,
            max_retries=max_retries,
        ), cfg.base_url
    except Exception as e:
        logger.warning("AI 去痕：OpenAI 客户端初始化失败：%s", e)
        return None, cfg.base_url


def clean_text(
    raw: Optional[str],
    *,
    use_ai: bool,
    client: Optional[OpenAI],
    model: str,
    force_local: bool = False,
) -> Tuple[str, str]:
    """
    Returns (cleaned_text, method) where method is 'api' or 'local'.
    force_local: 为 True 时仅用规则清理（如表格数据需保留制表符）。
    """
    if not raw:
        return "", "local"
    text = raw
    if len(text) > _MAX_INPUT_CHARS:
        text = text[:_MAX_INPUT_CHARS]
        logger.warning("AI 去痕：输入已截断至 %d 字", _MAX_INPUT_CHARS)

    if not force_local and use_ai and client:
        try:
            cleaned = clean_text_with_api(client, model, text)
            if cleaned:
                return sanitize_xml_compatible_text(cleaned), "api"
        except Exception as e:
            logger.error("AI 去痕 API 失败，回退本地规则：%s", e, exc_info=True)

    return sanitize_xml_compatible_text(clean_text_local(text)), "local"


_GENERATE_ARTICLE_SYS = """你是写作助手。用户会提出写作主题与要求。请用简体中文输出，结构必须规范：
1. 使用 Markdown 层级标题组织全文：主要章节用单独一行「## 章节名」；需要小节时用「### 小节名」。不要用一级「# 」。
2. 每个标题下用空行分隔多个正文段落；正文段落首行无需加井号，不要加粗或列表符号。
3. 不要输出「好的」「以下是」「综上所述」等套话开头或结尾，除非主题本身需要。
4. 不要代码块；表格类需求若用户明确要求再输出制表符表，否则用叙述性段落。"""

_GENERATE_TABLE_SYS = """用户会描述需要什么样的数据表格。
请只输出一张用制表符（Tab 键）分隔的表格：第一行为表头，后续每行一条数据；单元格内不要用制表符。
不要使用 Markdown、不要用竖线画表、不要用代码块，不要任何解释或前言。"""


def generate_content(
    client: OpenAI,
    model: str,
    user_prompt: str,
    kind: str = "article",
) -> str:
    """调用 ModelScope 兼容接口生成文章或表格（制表符分隔）。"""
    user_prompt = (user_prompt or "").strip()
    if not user_prompt:
        return ""
    sys_msg = _GENERATE_TABLE_SYS if kind == "table" else _GENERATE_ARTICLE_SYS
    max_tokens = int(os.getenv("AI_GENERATE_MAX_TOKENS", "8192"))
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": sys_msg},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.45,
        max_tokens=max_tokens,
    )
    return _extract_chat_completion_text(response)


_GENERATE_FILE_FILL_SYS = """你是文档编辑助手。用户会提供从 Word 或 Excel 抽取的纯文本（含段落与表格转写），并说明要如何修改或补全。
请严格依据「用户要求」处理「原文」：可补全空白、未完成句子、待填栏目；可修正语病与统一表述；保持专业、客观。
输出要求：
- 使用简体中文；用单独成行的 ## 章节名、### 小节名 组织层次（勿用一级 # ）；段落之间空一行。
- 直接输出正文，不要前言、「好的」「以下是修改后」等套话。
- 若用户明确要求保留表格数据，可用制表符分隔行列输出一张表；否则以叙述与分段为主。"""


def generate_file_fill(
    client: OpenAI,
    model: str,
    instructions: str,
    source_text: str,
    *,
    office_kind: str = "",
    docx_bytes: Optional[bytes] = None,
) -> Tuple[str, Optional[bytes]]:
    """
    根据说明对抽取自 Office 的正文做补全/修订。

    office_kind 为 docx 且传入 docx_bytes 时，**仅**走 Aspose.Words 结构化管线，返回 ("", .docx 字节)；
    不再使用 python-docx 抽取全文再走纯文本模型回写。
    """
    instructions = (instructions or "").strip()
    if not instructions:
        return "", None

    if (office_kind or "").strip().lower() == "docx" and docx_bytes:
        from utils.aspose_docx_processor import (  # noqa: PLC0415
            ensure_aspose_installed,
            run_aspose_structured_fill,
        )

        ensure_aspose_installed()
        out = run_aspose_structured_fill(
            docx_bytes,
            instructions,
            client,
            model=model,
        )
        return "", out

    source_text = (source_text or "").strip()
    if not source_text:
        return "", None

    max_in = int(os.getenv("AI_FILE_FILL_MAX_CHARS", "120000"))
    if len(source_text) > max_in:
        source_text = (
            source_text[:max_in]
            + "\n\n（原文过长已截断至此，请仅就以上可见部分结合用户要求输出完整稿。）"
        )
    user_block = f"【用户要求】\n{instructions}\n\n【原文】\n{source_text}"
    system = generate_file_fill_system_prompt(
        office_kind=office_kind,
        base_system=_GENERATE_FILE_FILL_SYS,
    )
    max_tokens = int(os.getenv("AI_FILE_FILL_MAX_TOKENS", "8192"))
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_block},
        ],
        temperature=0.35,
        max_tokens=max_tokens,
    )
    return _extract_chat_completion_text(response), None


def _detect_space_separated_credential_table(lines: list[str]) -> bool:
    """
    模型常输出「user001 acc001 Pass@...」这类空格分列、无制表符的示例表。
    至少 3 行匹配 user### + acc### 前缀时才启用，避免把普通英文句子拆成三列。
    """
    pat = re.compile(r"^user\d+\s+acc\d+\s+\S", re.IGNORECASE)
    n = sum(1 for ln in lines if pat.match(ln.strip()))
    return n >= max(3, min(5, len(lines) // 2))


def parse_table_lines(text: str) -> list[list[str]]:
    """将模型输出的制表符、竖线、多空格或「用户/账号/密码」式空格分列解析为二维表（至少一行）。"""
    lines = [ln.rstrip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return [[(text or "").strip()]]
    credential_like = _detect_space_separated_credential_table(lines)
    rows: list[list[str]] = []
    triple = re.compile(r"^(\S+)\s+(\S+)\s+(.+)$")
    for ln in lines:
        s = ln.strip()
        if "\t" in ln:
            rows.append([c.strip() for c in ln.split("\t")])
        elif "|" in ln and ln.count("|") >= 1:
            parts = [c.strip() for c in ln.split("|")]
            parts = [p for p in parts if p]
            if parts:
                rows.append(parts)
            else:
                rows.append([s])
        elif credential_like and (m := triple.match(s)):
            rows.append([m.group(1), m.group(2), m.group(3)])
        elif re.search(r" {2,}", s):
            rows.append([p.strip() for p in re.split(r" {2,}", s) if p.strip()])
        else:
            rows.append([s])
    return rows


def export_optional_deps_status() -> dict:
    """供前端提示：是否已安装导出依赖。"""
    out: dict = {}
    for key, mod in (
        ("python_docx", "docx"),
        ("openpyxl", "openpyxl"),
        ("reportlab", "reportlab"),
    ):
        try:
            __import__(mod)
            out[key] = True
        except ImportError:
            out[key] = False
    return out


def _require_docx():
    try:
        import docx

        return docx
    except ImportError as e:
        raise RuntimeError(
            "未安装 python-docx，请在项目目录执行: pip install python-docx"
        ) from e


def _require_openpyxl():
    try:
        import openpyxl

        return openpyxl
    except ImportError as e:
        raise RuntimeError(
            "未安装 openpyxl，请在项目目录执行: pip install openpyxl"
        ) from e


def _require_reportlab():
    try:
        import reportlab  # noqa: F401

        return True
    except ImportError as e:
        raise RuntimeError(
            "未安装 reportlab，请在项目目录执行: pip install reportlab"
        ) from e


def _find_cjk_font_path() -> Optional[str]:
    env_font = os.getenv("AI_CLEAN_PDF_FONT")
    candidates = [
        env_font,
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ]
    for p in candidates:
        if p and os.path.isfile(p):
            return p
    return None


def write_pdf_bytes(title: str, body: str) -> bytes:
    _patch_hashlib_md5_for_reportlab()
    _require_reportlab()
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    font_name = "Helvetica"
    font_path = _find_cjk_font_path()
    if font_path:
        try:
            if font_path.lower().endswith(".ttc"):
                pdfmetrics.registerFont(TTFont("CJK", font_path, subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
            font_name = "CJK"
        except Exception as e:
            logger.warning("注册 PDF 中文字体失败，英文可用：%s", e)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "T",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        leading=22,
    )
    body_style = ParagraphStyle(
        "TcdBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=20,
        firstLineIndent=22,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
    )
    h_styles = {
        1: ParagraphStyle(
            "TcdH1",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=15,
            leading=22,
            spaceBefore=12,
            spaceAfter=8,
        ),
        2: ParagraphStyle(
            "TcdH2",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=13.5,
            leading=19,
            spaceBefore=10,
            spaceAfter=6,
        ),
        3: ParagraphStyle(
            "TcdH3",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=12,
            leading=17,
            spaceBefore=8,
            spaceAfter=4,
        ),
    }

    def esc(s: str) -> str:
        s = sanitize_xml_compatible_text(s or "")
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    story = []
    story.append(Paragraph(esc(title or "文稿"), title_style))
    story.append(Spacer(1, 0.5 * cm))
    blocks = list(iter_document_blocks(body or ""))
    if not blocks and (body or "").strip():
        story.append(Paragraph(esc((body or "").strip()), body_style))
    else:
        for kind, level, content in blocks:
            if kind == "h" and level is not None:
                story.append(Paragraph(esc(content), h_styles.get(level, h_styles[3])))
            else:
                story.append(Paragraph(esc(content), body_style))
    doc.build(story)
    return buf.getvalue()


def write_docx_bytes(title: str, body: str) -> bytes:
    docx = _require_docx()
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Cm

    title = sanitize_xml_compatible_text(title or "") or "文稿"
    body = sanitize_xml_compatible_text(body or "")
    d = docx.Document()
    d.add_heading(title, level=1)
    blocks = list(iter_document_blocks(body or ""))
    if not blocks and (body or "").strip():
        p = d.add_paragraph((body or "").strip())
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        for kind, level, content in blocks:
            if kind == "h" and level is not None:
                d.add_heading(content, level=min(level + 1, 9))
            else:
                p = d.add_paragraph(content)
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = 1.25
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def write_xlsx_bytes(title: str, body: str) -> bytes:
    openpyxl = _require_openpyxl()
    from openpyxl.styles import Alignment, Font

    title = sanitize_xml_compatible_text(title or "") or "文稿"
    body = sanitize_xml_compatible_text(body or "")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "文稿"
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = body
    for row in ws.iter_rows(min_row=1, max_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.column_dimensions["A"].width = 100
    ws.row_dimensions[2].height = 280
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def write_xlsx_table_bytes(doc_title: str, rows: list[list[str]]) -> bytes:
    """将二维表写入 Excel：第 1 行文档标题，第 2 行起为表（首行为表头）。"""
    openpyxl = _require_openpyxl()
    from openpyxl.styles import Alignment, Font

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "表格"
    ws["A1"] = sanitize_xml_compatible_text(doc_title or "") or "表格"
    ws["A1"].font = Font(bold=True, size=14)
    start = 2
    if not rows:
        ws["A2"] = ""
    else:
        for ri, row in enumerate(rows, start=start):
            for ci, val in enumerate(row, start=1):
                v = sanitize_xml_compatible_text(val) if val is not None else ""
                c = ws.cell(row=ri, column=ci, value=v)
                c.alignment = Alignment(wrap_text=True, vertical="top")
                if ri == start:
                    c.font = Font(bold=True)
    from openpyxl.utils import get_column_letter

    max_cols = max((len(r) for r in rows), default=1) if rows else 1
    for i in range(1, max_cols + 1):
        ws.column_dimensions[get_column_letter(i)].width = 18
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def write_xlsx_smart_bytes(title: str, body: str) -> bytes:
    """
    若正文可解析为多列表格（制表符、竖线、多空格、user/acc 示例行等），则写入多列；
    否则与 write_xlsx_bytes 相同，整段置于 A2。
    """
    rows = parse_table_lines(body)
    max_cols = max((len(r) for r in rows), default=1)
    if max_cols >= 2:
        return write_xlsx_table_bytes(title, rows)
    return write_xlsx_bytes(title, body)
